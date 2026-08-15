from io import BytesIO
from unittest.mock import patch
from zipfile import BadZipFile

import pytest
from docx import Document

from parsing.extractors.docx import (
    DocxExtractionError,
    extract_docx_text,
)


def build_docx_bytes(document) -> bytes:
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def test_extract_docx_text_returns_nonempty_paragraphs():
    document = Document()

    document.add_paragraph("Ericka James")
    document.add_paragraph("")
    document.add_paragraph("Software Engineer")

    file_bytes = build_docx_bytes(document)

    text = extract_docx_text(file_bytes)

    assert text == (
        "Ericka James\n"
        "Software Engineer"
    )


def test_extract_docx_text_strips_paragraph_whitespace():
    document = Document()

    document.add_paragraph(
        "  Python Developer  "
    )

    file_bytes = build_docx_bytes(document)

    text = extract_docx_text(file_bytes)

    assert text == "Python Developer"


def test_extract_docx_text_preserves_bullet_structure():
    document = Document()

    document.add_paragraph(
        "Built an automation workflow.",
        style="List Bullet",
    )

    file_bytes = build_docx_bytes(document)

    text = extract_docx_text(file_bytes)

    assert text.startswith("• ")
    assert text.endswith(
        "Built an automation workflow."
    )


def test_extract_docx_text_extracts_table_rows():
    document = Document()

    document.add_paragraph("SKILLS")

    table = document.add_table(
        rows=2,
        cols=2,
    )

    table.cell(0, 0).text = "Languages"
    table.cell(0, 1).text = "Python, Java"

    table.cell(1, 0).text = "Cloud"
    table.cell(1, 1).text = "AWS"

    file_bytes = build_docx_bytes(document)

    text = extract_docx_text(file_bytes)

    assert "SKILLS" in text
    assert "Languages | Python, Java" in text
    assert "Cloud | AWS" in text


def test_extract_docx_text_preserves_document_order():
    document = Document()

    document.add_paragraph("Before table")

    table = document.add_table(
        rows=1,
        cols=2,
    )

    table.cell(0, 0).text = "Python"
    table.cell(0, 1).text = "Java"

    document.add_paragraph("After table")

    file_bytes = build_docx_bytes(document)

    text = extract_docx_text(file_bytes)

    lines = text.splitlines()

    assert lines == [
        "Before table",
        "Python | Java",
        "After table",
    ]


def test_extract_docx_text_rejects_empty_file():
    with pytest.raises(
        DocxExtractionError,
        match="DOCX file is empty",
    ):
        extract_docx_text(b"")


def test_extract_docx_text_rejects_malformed_document():
    with patch(
        "parsing.extractors.docx.Document",
        side_effect=BadZipFile("broken DOCX"),
    ):
        with pytest.raises(
            DocxExtractionError,
            match="DOCX file could not be read",
        ):
            extract_docx_text(b"broken")


def test_extract_docx_text_rejects_document_without_text():
    document = Document()

    file_bytes = build_docx_bytes(document)

    with pytest.raises(
        DocxExtractionError,
        match="No usable text",
    ):
        extract_docx_text(file_bytes)